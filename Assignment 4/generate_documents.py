import os
import pandas as pd
import streamlit as st
from fpdf import FPDF
from docx import Document


class DocumentGenerator:
    OUTPUT_DIR = "output"

    def __init__(self):
        os.makedirs(self.OUTPUT_DIR, exist_ok=True)

    def generate_pdf(self, employee_name, email, company, position, joining_date):
        """Generates PDF and Returns the path of the generated PDF"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            details = [
                "Employee Details",
                f"Employee Name: {employee_name}",
                f"Email: {email}",
                f"Company: {company}",
                f"Position: {position}",
                f"Joining Date: {joining_date}"
            ]
            for detail in details:
                pdf.cell(200, 10, txt=detail, ln=True)

            output_path = os.path.join(
                self.OUTPUT_DIR, f"{employee_name}_{company}.pdf")
            pdf.output(output_path)
            return output_path
        except Exception as e:
            st.error(f"Error generating PDF for {employee_name}: {e}")
            return None

    def generate_word(self, employee_name, email, company, position, joining_date):
        """Generates word file and Returns the path of the generated Word document"""
        try:
            doc = Document()
            details = {
                "Employee Name": employee_name,
                "Email": email,
                "Company": company,
                "Position": position,
                "Joining Date": joining_date
            }
            doc.add_heading("Employee Details", level=1)
            for key, value in details.items():
                doc.add_paragraph(f"{key}: {value}")

            output_path = os.path.join(
                self.OUTPUT_DIR, f"{employee_name}_{company}.docx")
            doc.save(output_path)
            return output_path
        except Exception as e:
            st.error(f"Error generating Word document for {
                     employee_name}: {e}")
            return None

    def process_excel(self, file):
        """Generates the output and Returns a list of generated file paths"""
        try:
            df = pd.read_excel(file)
            generated_files = []
            for _, row in df.iterrows():
                pdf_path = self.generate_pdf(
                    row['Name'], row['Email'], row['Company Name'], row['Position'], row['Joining Date'])
                word_path = self.generate_word(
                    row['Name'], row['Email'], row['Company Name'], row['Position'], row['Joining Date'])
                if pdf_path and word_path:
                    generated_files.append((pdf_path, word_path))
            return generated_files
        except Exception as e:
            st.error(f"Error processing Excel file: {e}")
            return []


def main():
    st.title("Bulk Document Generator")
    uploaded_file = st.file_uploader("Upload Excel File", type=["xlsx"])
    generator = DocumentGenerator()

    if uploaded_file:
        st.write("Processing file...")
        generated_files = generator.process_excel(uploaded_file)

        if generated_files:
            st.success("Documents Generated Successfully!")
            for pdf, word in generated_files:
                st.download_button(f"Download {os.path.basename(pdf)}", open(
                    pdf, "rb"), file_name=os.path.basename(pdf))
                st.download_button(f"Download {os.path.basename(word)}", open(
                    word, "rb"), file_name=os.path.basename(word))
        else:
            st.error("No documents were generated. Please check the input file.")


if __name__ == "__main__":
    main()
